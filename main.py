import os
import ast
import inspect
import importlib.util
from typing import Dict, List, Optional
from pathlib import Path

# For Word document (optional)
try:
    from docx import Document
except ImportError:
    Document = None

def analyze_function(node: ast.FunctionDef) -> Dict:
    """Extract function details from AST node"""
    func_details = {
        'name': node.name,
        'args': [],
        'returns': None,
        'docstring': ast.get_docstring(node) or ''
    }

    # Get arguments
    for arg in node.args.args:
        arg_info = {'name': arg.arg}
        if arg.annotation:
            arg_info['type'] = ast.unparse(arg.annotation)
        func_details['args'].append(arg_info)

    # Get return type
    if node.returns:
        func_details['returns'] = ast.unparse(node.returns)

    return func_details

def process_file(file_path: str) -> Dict:
    """Process a single Python file"""
    result = {'modules': {}}
    with open(file_path, 'r', encoding='utf-8') as f:
        try:
            tree = ast.parse(f.read())
            for node in ast.walk(tree):
                if isinstance(node, ast.FunctionDef):
                    module_name = os.path.basename(file_path).replace('.py', '')
                    result['modules'].setdefault(module_name, []).append(
                        analyze_function(node)
                    )
        except Exception as e:
            print(f"Error parsing {file_path}: {str(e)}")
    return result

def scan_project(root_dir: str) -> Dict:
    """Scan all Python files in directory"""
    project_data = {}
    for dirpath, _, filenames in os.walk(root_dir):
        for filename in filenames:
            if filename.endswith('.py') and filename != os.path.basename(__file__):
                full_path = os.path.join(dirpath, filename)
                data = process_file(full_path)
                if data['modules']:
                    project_data[full_path] = data
    return project_data

def generate_report(data: Dict, output_format: str = 'txt') -> None:
    """Generate documentation file"""
    if output_format == 'docx' and Document:
        doc = Document()
        doc.add_heading('Project Documentation', 0)
    else:
        lines = []

    for file_path, modules in data.items():
        if output_format == 'docx':
            doc.add_heading(f'File: {file_path}', level=1)
        else:
            lines.append(f"\nFile: {file_path}\n{'-'*40}")

        for module, functions in modules['modules'].items():
            if output_format == 'docx':
                doc.add_heading(f'Module: {module}', level=2)
            else:
                lines.append(f"\nModule: {module}\n{'='*40}")

            for func in functions:
                if output_format == 'docx':
                    doc.add_heading(func['name'], level=3)
                    doc.add_paragraph(f"Arguments: {[arg['name'] for arg in func['args']]}")
                    doc.add_paragraph(f"Returns: {func['returns'] or 'None'}")
                    doc.add_paragraph(f"Description: {func['docstring']}")
                else:
                    lines.append(f"\nFunction: {func['name']}")
                    lines.append(f"Inputs: {[arg['name'] for arg in func['args']]}")
                    lines.append(f"Returns: {func['returns'] or 'None'}")
                    lines.append(f"Description: {func['docstring']}")
                    lines.append('-'*40)

    if output_format == 'docx' and Document:
        doc.save('project_documentation.docx')
        print("Word document generated: project_documentation.docx")
    else:
        with open('project_documentation.txt', 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        print("Text document generated: project_documentation.txt")

if __name__ == "__main__":
    project_root = os.path.dirname(os.path.abspath(__file__))
    project_data = scan_project(project_root)
    
    # Try to generate Word doc if possible
    if Document:
        generate_report(project_data, output_format='docx')
    else:
        generate_report(project_data)
        print("Install python-docx for Word document support")
